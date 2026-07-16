"""Corrige pendências específicas de Figura e Tabelas conforme Manual USP/Esalq:
  - Remove ponto final da legenda da Figura 1
  - Garante a Fonte da Figura 1 (caso esteja vazia)
  - Alinha cabeçalhos das colunas 2+ centralizados
  - Alinha conteúdo numérico à direita (Tabela 2)
  - Alinha conteúdo de texto (Tabela 1) - valores centralizado, justificativa justificado
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SRC = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (final).docx")
DST = SRC

# Números (com ou sem vírgula decimal, % ou outros)
NUM_RE = re.compile(r"^\d[\d.,%]*\s*(?:\(.+\))?$")


def set_paragraph_text(paragraph, text: str):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def align_cell(cell, alignment):
    for p in cell.paragraphs:
        p.alignment = alignment


def main():
    doc = Document(str(SRC))

    # 1. Remover ponto final da legenda da Figura 1
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Figura 1.") and t.endswith("."):
            # Remove só o ponto FINAL (o ponto após "Figura 1." permanece)
            # "Figura 1. Evolução da função de perda (Cross-Entropy) durante o treinamento."
            # → "Figura 1. Evolução da função de perda (Cross-Entropy) durante o treinamento"
            novo = t[:-1]
            set_paragraph_text(p, novo)
            print(f"[figura] ponto final removido: {novo[:80]!r}")
            # Verifica a fonte (próximo parágrafo)
            if i + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[i + 1]
                if not next_p.text.strip():
                    # Está vazio — substituir pela fonte
                    set_paragraph_text(next_p, "Fonte: Resultados originais da pesquisa")
                    print(f"[figura] fonte inserida no parágrafo vazio seguinte")
                elif next_p.text.strip().endswith("."):
                    # Remove ponto final da Fonte, se tiver
                    fonte_txt = next_p.text.rstrip(".")
                    set_paragraph_text(next_p, fonte_txt)
                    print(f"[figura] ponto final da fonte removido")
            break

    # Também normaliza Fonte das Tabelas sem ponto final
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Fonte:") and t.endswith("."):
            set_paragraph_text(p, t.rstrip("."))
            print(f"[fonte-tabela] ponto final removido: {t[:60]!r}")

    # 2. Alinhamento das tabelas
    for ti, tab in enumerate(doc.tables):
        n_cols = len(tab.columns)
        n_rows = len(tab.rows)
        print(f"\n[Tabela {ti+1}] {n_rows} linhas × {n_cols} colunas")

        for ri, row in enumerate(tab.rows):
            for ci, cell in enumerate(row.cells):
                # Linha 0 = cabeçalho
                if ri == 0:
                    if ci == 0:
                        align_cell(cell, WD_ALIGN_PARAGRAPH.LEFT)
                    else:
                        align_cell(cell, WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    # Corpo: primeira coluna à esquerda; demais centralizado ou direita (números)
                    if ci == 0:
                        align_cell(cell, WD_ALIGN_PARAGRAPH.LEFT)
                    else:
                        # Decide: se o conteúdo é numérico, vai à direita; senão centralizado
                        txt = cell.text.strip()
                        if NUM_RE.match(txt) or txt.replace(",", ".").replace("%", "").replace(" ", "").replace("Milhões","").replace("tokens","").strip().replace(".","").isdigit():
                            align_cell(cell, WD_ALIGN_PARAGRAPH.RIGHT)
                        else:
                            # Se tem texto longo (justificativa), justificado; senão centralizado
                            if len(txt) > 35:
                                align_cell(cell, WD_ALIGN_PARAGRAPH.JUSTIFY)
                            else:
                                align_cell(cell, WD_ALIGN_PARAGRAPH.CENTER)
        print(f"  alinhamentos aplicados: 1ª col esquerda, cabeçalhos demais centralizados, números à direita")

    doc.save(str(DST))
    print(f"\n=== SALVO: {DST} ===")


if __name__ == "__main__":
    main()
