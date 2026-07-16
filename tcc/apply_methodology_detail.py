"""Aplica os ajustes metodológicos pedidos pela orientadora (v4):
  1. Função de perda (cross-entropy, mascaramento -100)
  2. Estratégia de divisão treino/validação (90/10 sequencial)
  3. Critério formal de early stopping (delta 0.01, paciência 3)
  4. Configuração do batching (random sampling, batch=1, grad_accum=16)
  5. Definição operacional das métricas (loss e acurácia de token)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path("/tmp/tcc_v4.docx")
DST = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (template).docx")


# -----------------------------------------------------------
# Textos novos
# -----------------------------------------------------------
# Expansão da Otimização: inclui função de perda, batching, clipping
NEW_OTIMIZACAO = (
    "O treinamento foi realizado utilizando o otimizador AdamW (Loshchilov; Hutter, "
    "2019), que desacopla o decaimento de peso da otimização do gradiente, configurado "
    "com taxa de aprendizado inicial de 3e-4, decaimento de peso de 1e-2 e demais "
    "hiperparâmetros em valores padrão (β₁=0,9; β₂=0,999; ε=1e-8). A taxa de "
    "aprendizado segue agendamento por decaimento cosseno (CosineAnnealingLR) até o "
    "valor mínimo de 1e-5 ao longo das iterações totais. Empregou-se gradient clipping "
    "com norma máxima de 1,0 para estabilidade numérica do gradiente. Como técnica de "
    "regularização para prevenir o sobreajuste (overfitting) em cenários de dados "
    "limitados, aplicou-se Dropout (Srivastava et al., 2014) com probabilidade de 0,1 "
    "nas matrizes de atenção e nas camadas residuais, ativo apenas durante o treino e "
    "desativado na avaliação e geração."
)

# Nova subseção: Função de Perda e Configuração de Batches (entre Otimização e Reprodutibilidade)
# Vamos adicionar ANTES de Otimização pra ficar mais lógico
NEW_LOSS_TITLE = "Função de Perda e Configuração de Batches"

NEW_LOSS_PAR1 = (
    "A função objetivo é a entropia cruzada (cross-entropy) entre a distribuição "
    "prevista pelo modelo sobre o vocabulário e o token seguinte real, computada pela "
    "rotina torch.nn.functional.cross_entropy. Durante o pré-treino causal todos os "
    "tokens da sequência contribuem para a perda; durante o fine-tuning supervisionado "
    "(SFT) o cálculo é restrito aos tokens da resposta, mascarando-se os tokens do "
    "prompt com o valor −100, constante padrão do PyTorch para índices ignorados no "
    "cômputo da perda."
)

NEW_LOSS_PAR2 = (
    "Os batches são construídos por amostragem aleatória uniforme de posições iniciais "
    "no tensor tokenizado do corpus de treino, pareando sequências de janela fixa de "
    "256 tokens (block_size). Adotou-se batch_size efetivo de 16 sequências por passo "
    "de otimização, atingido pela combinação de batch_size físico de 1 com "
    "gradient accumulation de 16 passos, estratégia necessária para respeitar o limite "
    "de 6 GB de VRAM da GPU utilizada. A divisão do corpus tokenizado em conjuntos de "
    "treino e validação seguiu corte sequencial na proporção de 90% (treino) e 10% "
    "(validação), prática comum no pré-treino autorregressivo de modelos causais em "
    "corpus contínuo."
)

# Reprodutibilidade — com critério formal de early stopping
NEW_REPRO = (
    "Para permitir a reprodução do experimento descrito, registram-se os seguintes "
    "parâmetros operacionais. Os checkpoints V6 e V7 foram gerados com o conjunto de "
    "hiperparâmetros reportado na Tabela 1, utilizando a biblioteca PyTorch 2.11 sobre "
    "CUDA em GPU NVIDIA GeForce RTX 3050 Laptop (6 GB de VRAM), CPU Intel Core "
    "i5-13420H e 7 GB de memória RAM. O split entre treino e validação seguiu "
    "proporção de 90%/10% por corte sequencial (ver subseção anterior). O experimento "
    "não utilizou seed fixa para a amostragem aleatória de batches nem para a "
    "inicialização dos pesos, o que pode introduzir variação entre execuções; "
    "recomenda-se, em replicações futuras, a fixação explícita de seed para maior "
    "determinismo. O critério formal de parada (early stopping) foi definido como a "
    "ausência de melhoria mínima de 0,01 na perda média de validação (calculada sobre "
    "50 batches aleatórios) em três avaliações consecutivas, sendo cada avaliação "
    "realizada a cada 200 iterações. Registraram-se o V6 ao final da iteração 1.400 e "
    "o V7 ao final da iteração 5.000."
)

# Nova subseção: Métricas de Avaliação (ANTES de "Resultados e Discussão")
NEW_METRICS_TITLE = "Métricas de Avaliação"

NEW_METRICS_INTRO = (
    "Duas métricas quantitativas foram empregadas para monitorar o progresso do "
    "treinamento, ambas computadas sobre 50 batches aleatórios do conjunto de "
    "validação com o modelo em modo de avaliação (dropout desativado), a cada 200 "
    "iterações. A avaliação qualitativa complementar é realizada por inspeção de "
    "gerações a partir de prompts em português, com temperatura de amostragem 0,7 e "
    "truncamento top-k de 50, analisando as saídas quanto à fluência, estrutura "
    "sintática e coerência local."
)

NEW_METRICS_LOSS = (
    "(i) Perda de validação (loss): média aritmética da entropia cruzada por token "
    "calculada entre a distribuição prevista e o token real. Valores menores indicam "
    "maior probabilidade atribuída pelo modelo ao token real, refletindo melhor "
    "capacidade preditiva."
)

NEW_METRICS_ACC = (
    "(ii) Acurácia de predição de token (token-level accuracy): fração de tokens para "
    "os quais a posição com maior probabilidade na saída do modelo (argmax sobre os "
    "logits) coincide com o token real da sequência. Valores mais próximos de 1 "
    "indicam maior concordância determinística entre predições e rótulos."
)


# -----------------------------------------------------------
# Helpers
# -----------------------------------------------------------
def set_paragraph_text_as_new(paragraph, text: str):
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_after(ref_paragraph, text: str = "", bold: bool = False, indent: bool = True):
    """Insere parágrafo logo após o referenciado, com Arial 11."""
    new_elem = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_elem)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_elem, ref_paragraph._parent)
    if text:
        run = new_para.add_run(text)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(11)
    if indent:
        new_para.paragraph_format.first_line_indent = Cm(1.25)
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return new_para


# -----------------------------------------------------------
# Main
# -----------------------------------------------------------
def main():
    doc = Document(str(SRC))
    print(f"[carregado] {len(doc.paragraphs)} parágrafos\n")

    # Localizadores
    def find_para(substr_or_exact, exact=False, after=0):
        for i in range(after, len(doc.paragraphs)):
            t = doc.paragraphs[i].text.strip()
            if exact and t == substr_or_exact:
                return i
            if not exact and substr_or_exact in t:
                return i
        return None

    # 1. Expande "Otimização e Regularização"
    i_otim = find_para("utilizando o otimizador AdamW")
    if i_otim is not None:
        set_paragraph_text_as_new(doc.paragraphs[i_otim], NEW_OTIMIZACAO)
        doc.paragraphs[i_otim].paragraph_format.first_line_indent = Cm(1.25)
        doc.paragraphs[i_otim].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        print(f"[ok 1] par {i_otim} (Otimização) expandido com função de perda/batching/clipping")

    # 2. Adicionar subseção "Função de Perda e Configuração de Batches"
    #    ANTES de "Otimização e Regularização"
    i_otim_title = find_para("Otimização e Regularização", exact=True)
    if i_otim_title is not None:
        # Inserir título + 2 parágrafos ANTES do título de Otimização
        title_para = doc.paragraphs[i_otim_title]
        # Adiciona um parágrafo vazio antes (espaçamento)
        title_para.insert_paragraph_before("")
        # Insere título da nova subseção (negrito)
        new_title = title_para.insert_paragraph_before(NEW_LOSS_TITLE)
        new_title.paragraph_format.first_line_indent = Cm(0)  # subtítulo sem indent
        new_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in new_title.runs:
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(11)
        # Par vazio de separação
        title_para.insert_paragraph_before("")
        # Par 1
        p1 = title_para.insert_paragraph_before(NEW_LOSS_PAR1)
        p1.paragraph_format.first_line_indent = Cm(1.25)
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p1.runs:
            r.font.name = "Arial"; r.font.size = Pt(11)
        # Par 2
        p2 = title_para.insert_paragraph_before(NEW_LOSS_PAR2)
        p2.paragraph_format.first_line_indent = Cm(1.25)
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p2.runs:
            r.font.name = "Arial"; r.font.size = Pt(11)
        # Par vazio separando da próxima subseção
        title_para.insert_paragraph_before("")
        print(f"[ok 2] nova subseção 'Função de Perda e Configuração de Batches' inserida antes de Otimização")

    # 3. Atualiza Reprodutibilidade com critério formal de early stopping
    i_rep = find_para("reprodução do experimento")
    if i_rep is not None:
        set_paragraph_text_as_new(doc.paragraphs[i_rep], NEW_REPRO)
        doc.paragraphs[i_rep].paragraph_format.first_line_indent = Cm(1.25)
        doc.paragraphs[i_rep].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        print(f"[ok 3] par {i_rep} (Reprodutibilidade) atualizado com early stopping formal")

    # 4. Adicionar subseção "Métricas de Avaliação" ANTES de "Resultados e Discussão"
    i_res = find_para("Resultados e Discussão", exact=True)
    if i_res is not None:
        res_para = doc.paragraphs[i_res]
        # Par vazio pra separar
        res_para.insert_paragraph_before("")
        # Título da subseção
        title_metrics = res_para.insert_paragraph_before(NEW_METRICS_TITLE)
        title_metrics.paragraph_format.first_line_indent = Cm(0)
        title_metrics.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in title_metrics.runs:
            r.bold = True; r.font.name = "Arial"; r.font.size = Pt(11)
        res_para.insert_paragraph_before("")
        # Intro
        p_intro = res_para.insert_paragraph_before(NEW_METRICS_INTRO)
        p_intro.paragraph_format.first_line_indent = Cm(1.25)
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p_intro.runs:
            r.font.name = "Arial"; r.font.size = Pt(11)
        # Loss
        p_loss = res_para.insert_paragraph_before(NEW_METRICS_LOSS)
        p_loss.paragraph_format.first_line_indent = Cm(1.25)
        p_loss.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p_loss.runs:
            r.font.name = "Arial"; r.font.size = Pt(11)
        # Acurácia
        p_acc = res_para.insert_paragraph_before(NEW_METRICS_ACC)
        p_acc.paragraph_format.first_line_indent = Cm(1.25)
        p_acc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p_acc.runs:
            r.font.name = "Arial"; r.font.size = Pt(11)
        # Par vazio de transição pra Resultados
        res_para.insert_paragraph_before("")
        print(f"[ok 4] nova subseção 'Métricas de Avaliação' inserida antes de 'Resultados e Discussão'")

    doc.save(str(DST))
    print(f"\n=== SALVO: {DST} ===")


if __name__ == "__main__":
    main()
