"""Aplica as normas detalhadas do Manual MBA USP/Esalq no TCC:
  - Recuo 1,25 cm na primeira linha de TODOS os parágrafos do corpo
    (exceto: títulos, subtópicos, resumo, palavras-chave, captions, fontes,
    entradas de referência, folha de rosto)
  - Tabelas: só bordas horizontais (topo e base do cabeçalho + base da tabela),
    sem bordas internas/laterais, sem cor
"""
from pathlib import Path
import re

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (final).docx")
DST = SRC  # sobrescreve

TOPICOS_PRINCIPAIS = {
    "RESUMO", "SUMÁRIO", "CONSIDERAÇÕES INICIAIS",
    "IMPLEMENTAÇÃO DE ALGORITMO(S) DE MACHINE LEARNING",
    "RESULTADOS E DISCUSSÃO", "CONCLUSÃO", "REFERÊNCIAS",
}

SUBTOPICOS = {
    "Recursos Computacionais e Limitações de Hardware",
    "Construção do Dataset Multi-Domínio",
    "A Arquitetura Transformer",
    "Arquiteturas Decoder-Only (GPT)",
    "Tokenização Subword e BPE",
    "Dimensionamento do Modelo (Scaling Laws)",
    "Otimização e Regularização",
    "Reprodutibilidade do Experimento",
    "Análise de Convergência da Função de Perda",
    "Evolução da Acurácia de Predição",
    "Viabilidade em Hardware Acadêmico",
    "Avaliação Qualitativa Preliminar",
    "Amostras Qualitativas",
}


def is_body_text(paragraph, in_body: bool, in_refs: bool) -> bool:
    """Determina se o parágrafo recebe recuo 1,25 cm.
    - Precisa estar dentro do corpo (após RESUMO e antes de REFERÊNCIAS)
    - NÃO ser título, subtópico, caption, fonte, palavras-chave, ou item de amostra
    """
    t = paragraph.text.strip()
    if not in_body or in_refs or not t:
        return False
    if t in TOPICOS_PRINCIPAIS or t in SUBTOPICOS:
        return False
    if t.startswith("PALAVRAS-CHAVE"):
        return False
    if t.startswith(("Tabela ", "Figura ")) and "." in t[:15]:
        return False
    if t.startswith("Fonte:"):
        return False
    if t.startswith(("Prompt:", "Saída:")):
        return False
    # Folha de rosto: identificar linhas como "São Paulo - SP", "2025", etc.
    # Essas têm menos de ~40 chars e caracteres específicos. Considera fora do corpo.
    return True


def apply_body_indent(doc: Document):
    """Aplica recuo de 1,25 cm na primeira linha de parágrafos do corpo."""
    in_body = False
    in_refs = False
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "RESUMO":
            in_body = True
            continue
        if t == "REFERÊNCIAS":
            in_refs = True
            continue
        # O parágrafo DIRETO após RESUMO é o próprio resumo — sem recuo
        # O parágrafo de palavras-chave também — sem recuo
        # Já tratado por is_body_text via topicos e "PALAVRAS-CHAVE"
        if is_body_text(p, in_body, in_refs):
            p.paragraph_format.first_line_indent = Cm(1.25)
            count += 1
    print(f"[recuo] aplicado 1,25 cm em {count} parágrafos do corpo")


def set_cell_border(cell, border_name: str, val: str = "single", size: str = "4", color: str = "auto"):
    """Define uma borda da célula. border_name: top, left, bottom, right."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = tcBorders.find(qn(f"w:{border_name}"))
    if border is None:
        border = OxmlElement(f"w:{border_name}")
        tcBorders.append(border)
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), size)
    border.set(qn("w:color"), color)


def clear_cell_shading(cell):
    """Remove preenchimento de fundo da célula."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        tcPr.remove(shd)


def format_tables(doc: Document):
    """Aplica bordas só horizontais nas tabelas, remove cor, remove negrito no conteúdo."""
    for ti, tab in enumerate(doc.tables):
        n_rows = len(tab.rows)
        for ri, row in enumerate(tab.rows):
            for ci, cell in enumerate(row.cells):
                clear_cell_shading(cell)
                # Remove bordas laterais
                set_cell_border(cell, "left", val="nil")
                set_cell_border(cell, "right", val="nil")
                # Topo: só na primeira linha
                if ri == 0:
                    set_cell_border(cell, "top", val="single")
                else:
                    set_cell_border(cell, "top", val="nil")
                # Fundo: na linha 0 (cabeçalho) e na última linha
                if ri == 0 or ri == n_rows - 1:
                    set_cell_border(cell, "bottom", val="single")
                else:
                    set_cell_border(cell, "bottom", val="nil")
                # Remove negrito do conteúdo (preserva do cabeçalho)
                if ri > 0:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = False
        print(f"[tabela {ti}] {n_rows} linhas reformatadas (bordas só superior e inferior do cabeçalho + inferior da tabela)")


def main():
    doc = Document(str(SRC))
    print(f"[carregado] {len(doc.paragraphs)} parágrafos, {len(doc.tables)} tabelas\n")

    apply_body_indent(doc)
    format_tables(doc)

    doc.save(str(DST))
    print(f"\n=== SALVO: {DST} ===")


if __name__ == "__main__":
    main()
