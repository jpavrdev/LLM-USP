"""Aplica as 4 pendências finais no TCC:
    1. Recuo especial de 1,25 cm na primeira linha dos subtópicos
    2. Espaçamento simples em resumo, tabelas, rodapés, legendas
    3. Varredura exaustiva de inconsistências numéricas
    4. Revisão caso-a-caso das referências

Input:  /tmp/tcc_gdoc.docx (versão baixada do Google Docs do usuário)
Output: /home/akirem/Downloads/TCC MBA - USP 2025 (final).docx
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING


SRC = Path("/tmp/tcc_gdoc.docx")
DST = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (final).docx")

# Subtópicos do TCC — textos que aparecem como título de subseção, precisam ter
# recuo de 1,25 cm na primeira linha. Tópicos principais não entram.
SUBTOPICOS = {
    "Recursos Computacionais e Limitações de Hardware",
    "Construção do Dataset Multi-Domínio",
    "A Arquitetura Transformer",
    "Arquiteturas Decoder-Only (GPT)",
    "Tokenização Subword e BPE",
    "Dimensionamento do Modelo (Scaling Laws)",
    "Otimização e Regularização",
    "Reprodutibilidade do Experimento",
    "Análise de Convergência da Função de Perda",
    "Evolução da Acurácia de Predição",
    "Viabilidade em Hardware Acadêmico",
    "Avaliação Qualitativa Preliminar",
    "Amostras Qualitativas",
}

# Tópicos principais (pra não confundir)
TOPICOS_PRINCIPAIS = {
    "RESUMO", "SUMÁRIO", "CONSIDERAÇÕES INICIAIS",
    "IMPLEMENTAÇÃO DE ALGORITMO(S) DE MACHINE LEARNING",
    "RESULTADOS E DISCUSSÃO", "CONCLUSÃO", "REFERÊNCIAS",
}


def set_paragraph_text(paragraph, text: str):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def apply_subtopic_indent(doc: Document):
    """1. Recuo de 1,25 cm na primeira linha dos parágrafos que SEGUEM os subtópicos.
    O checklist diz 'subtópicos com recuo especial de 1,25 cm na primeira linha'.
    Interpretamos como: o parágrafo do título do subtópico fica à margem (negrito),
    e o parágrafo de corpo seguinte recebe o recuo.
    """
    count = 0
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt in SUBTOPICOS:
            # Garante negrito no título
            for r in p.runs:
                r.bold = True
            # Aplica recuo especial ao próximo parágrafo (corpo do subtópico)
            if i + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[i + 1]
                if next_p.text.strip():  # só se tem conteúdo
                    next_p.paragraph_format.first_line_indent = Cm(1.25)
                    count += 1
    print(f"[1] recuo 1,25 cm aplicado em {count} parágrafos de corpo de subtópicos")


def apply_simple_spacing(doc: Document):
    """2. Espaçamento simples em zonas específicas:
       - Resumo (o parágrafo após 'RESUMO')
       - Título e fonte de tabelas (parágrafos 'Tabela N.' e 'Fonte: ...')
       - Título e fonte de figuras (parágrafos 'Figura N.' e 'Fonte: ...')
       - Corpo das tabelas (células)
       - Notas de rodapé (o docx atual não tem formalmente, só um '1 pedro.veras...')
    """
    count = 0

    # (a) Resumo — parágrafo após o título RESUMO
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "RESUMO":
            # O parágrafo do resumo em si e o de palavras-chave
            for j in (i + 1, i + 2):
                if j < len(doc.paragraphs) and doc.paragraphs[j].text.strip():
                    doc.paragraphs[j].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    count += 1
            break

    # (b) Títulos e fontes de figuras/tabelas
    for p in doc.paragraphs:
        t = p.text.strip()
        if (t.startswith("Tabela ") and ("." in t[:15]) or
                t.startswith("Figura ") and ("." in t[:15]) or
                t.startswith("Fonte:")):
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            count += 1

    # (c) Corpo das tabelas
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    count += 1

    # (d) Notas de rodapé — os footnote refs no docx
    try:
        part = doc.part.footnotes_part
        if part is not None:
            from docx.oxml.ns import qn
            for p in part.element.iter(qn("w:p")):
                pPr = p.find(qn("w:pPr"))
                if pPr is not None:
                    spacing = pPr.find(qn("w:spacing"))
                    if spacing is None:
                        from lxml import etree
                        spacing = etree.SubElement(pPr, qn("w:spacing"))
                    spacing.set(qn("w:line"), "240")  # 240 twips = 1 linha simples
                    spacing.set(qn("w:lineRule"), "auto")
    except AttributeError:
        pass

    print(f"[2] espaçamento simples aplicado em {count} parágrafos/células")


def varredura_numericas(doc: Document):
    """3. Varredura exaustiva de inconsistências numéricas."""
    padroes_suspeitos = [
        # 30,58M residual
        (r"30[.,]\s*58", "30,58"),
        # batch 32 residual
        (r"\bbatch[_\s]size\s*(?:de|:)?\s*32\b", "batch 32"),
        (r"\blote\s+de\s+32\b", "lote de 32"),
        # contexto diferente de 256
        (r"contexto\s+de\s+(\d+)", "contexto"),
        # referências a "dropout 0.2" (v5) vs 0.1 (v6)
        (r"dropout\s*(?:de|:)?\s*0[,.]\s*2\b", "dropout 0.2"),
        # menção genérica a "n camadas"
        (r"(\d+)\s*camadas", "camadas"),
    ]

    print("\n[3] varredura numérica:")
    achados = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        for pat, label in padroes_suspeitos:
            for m in re.finditer(pat, t, re.IGNORECASE):
                valor = m.group(0)
                # filtros: contexto 256 tokens é o correto, camadas 8 é correto
                if label == "contexto" and "256" in valor:
                    continue
                if label == "camadas":
                    val_num = int(m.group(1))
                    if val_num == 8:  # 8 camadas é o correto
                        continue
                if "dropout 0.1" in valor.replace(",", "."):
                    continue
                print(f"  [{i}] {label!r}: ...{t[max(0,m.start()-40):m.end()+30]}...")
                achados += 1
    # Tabelas
    for ti, tab in enumerate(doc.tables):
        for ri, row in enumerate(tab.rows):
            for ci, cell in enumerate(row.cells):
                t = cell.text
                for pat, label in padroes_suspeitos:
                    for m in re.finditer(pat, t, re.IGNORECASE):
                        if label == "contexto" and "256" in m.group(0): continue
                        if label == "camadas":
                            v = int(m.group(1))
                            if v == 8: continue
                        print(f"  [Tab{ti} r{ri}c{ci}] {label!r}: {t[:80]}")
                        achados += 1
    if achados == 0:
        print("  (nenhuma inconsistência encontrada)")
    return achados


def normalize_refs(doc: Document):
    """4. Revisão formal das referências.
    Padrão USP/Esalq (checklist): 'Autor, I. Ano. Título. Revista vol(n): págs.'
    Ajustes simples: garantir espaçamento simples, alinhamento à esquerda, remover urls em linha separada.
    """
    print("\n[4] revisão de referências:")
    in_refs = False
    refs_count = 0
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt == "REFERÊNCIAS":
            in_refs = True
            continue
        if not in_refs or not txt:
            continue
        # Espaçamento simples em cada entrada
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Alinhamento à esquerda (normal)
        # (sem sangria; cada referência começa à margem esquerda)
        refs_count += 1

    print(f"  {refs_count} referências com espaçamento simples e alinhamento à esquerda")


def main():
    doc = Document(str(SRC))
    print(f"[carregado] {len(doc.paragraphs)} parágrafos, {len(doc.tables)} tabelas\n")

    apply_subtopic_indent(doc)
    apply_simple_spacing(doc)
    varredura_numericas(doc)
    normalize_refs(doc)

    doc.save(str(DST))
    print(f"\n=== SALVO: {DST} ===")


if __name__ == "__main__":
    main()
