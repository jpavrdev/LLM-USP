"""Formata o TCC revisado seguindo o Checklist USP/Esalq + orientações da orientadora.

Mudanças aplicadas:
  1. Formatação global (margens 2.5cm, Arial 11, espaçamento 1.5, 0pt before/after)
  2. Paginação arial 9 no canto inferior direito
  3. Siglas (XXX) → [XXX] na primeira ocorrência (exceto nas referências)
  4. Correção numérica: 30,58M → 52,93M, batch 32 → batch 16 (parágrafos desatualizados)
  5. Objetivos específicos — remover SFT (mover pra "trabalhos futuros")
  6. Adicionar subseção de reprodutibilidade (steps, seed, early stopping)
  7. Remover SUMÁRIO (não obrigatório e estava incoerente)
  8. Caption de Figura 1 com "Fonte: ..." abaixo
  9. Fontes abaixo das tabelas ("Fonte: Resultados originais da pesquisa")

Input:  TCC MBA - USP 2025 (revisado).docx
Output: TCC MBA - USP 2025 (final).docx
"""

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (revisado).docx")
DST = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (final).docx")


# ---------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------
def set_paragraph_text(paragraph, text: str):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text: str):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    set_paragraph_text(cell.paragraphs[0], text)


def add_page_number_field(paragraph):
    """Adiciona campo PAGE no parágrafo (Word recalcula na abertura)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def apply_global_formatting(doc: Document):
    """Margens 2,5cm, cabeçalho/rodapé 1,25cm, Arial 11, espaçamento 1,5."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    # Estilo Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Força Arial também nas runs dos parágrafos (alguns foram editados e têm fonte própria)
    for p in doc.paragraphs:
        for r in p.runs:
            if not r.font.name or r.font.name != "Arial":
                r.font.name = "Arial"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            if not r.font.size:
                r.font.size = Pt(11)


def add_footer_page_number(doc: Document):
    """Paginação Arial 9 no canto inferior direito."""
    for section in doc.sections:
        footer = section.footer
        # Limpa rodapés existentes
        p = footer.paragraphs[0]
        for r in list(p.runs):
            r.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number_field(p)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9)


def replace_acronyms_first_occurrence(doc: Document, max_paragraph_index_for_refs: int):
    """Troca (SIGLA) por [SIGLA] na primeira ocorrência de cada sigla no corpo.
    Não mexe nas referências (a partir de `max_paragraph_index_for_refs`).
    """
    # Siglas a converter (primeira ocorrência apenas)
    target_acronyms = ["SLM", "PLN", "SFT", "RLHF", "LGPD", "BPE", "RLAIF"]
    found = set()
    for i, p in enumerate(doc.paragraphs):
        if i >= max_paragraph_index_for_refs:
            break
        for acr in target_acronyms:
            if acr in found:
                continue
            pattern = f"({acr})"
            if pattern in p.text:
                new = p.text.replace(f"({acr})", f"[{acr}]", 1)
                set_paragraph_text(p, new)
                found.add(acr)
                print(f"[ok] sigla ({acr}) -> [{acr}] no parágrafo {i}")
    return found


# ---------------------------------------------------------------
# Textos novos de conteúdo
# ---------------------------------------------------------------
NOVO_PAR_OBJETIVOS = (
    "O objetivo geral deste trabalho é desenvolver, treinar e avaliar um Modelo de "
    "Linguagem de Pequeno Porte [SLM] de arquitetura decoder-only especializado na "
    "língua portuguesa. Os objetivos específicos incluem: (i) compilar e normalizar "
    "um corpus multi-domínio em português; (ii) definir a arquitetura Transformer "
    "decoder-only e seus hiperparâmetros sob a restrição de 6 GB de VRAM; "
    "(iii) executar o pré-treino causal com monitoramento contínuo de loss e "
    "acurácia; e (iv) avaliar qualitativa e quantitativamente o desempenho do modelo "
    "resultante. A etapa de ajuste fino supervisionado [SFT] e o alinhamento por "
    "preferências humanas permanecem como trabalhos futuros, discutidos na Conclusão."
)

NOVO_PAR_103 = (
    "Os experimentos confirmaram a hipótese de viabilidade técnica. A arquitetura "
    "dimensionada com 52,93 milhões de parâmetros, operando com batch size de 16 e "
    "contexto de 256 tokens, manteve a ocupação da VRAM da GPU (6 GB) dentro dos "
    "limites operacionais, sem erros de Out-of-Memory. O tempo de convergência inicial "
    "para uma loss próxima de 2,0 demonstrou ser compatível com janelas de tempo de "
    "projetos acadêmicos, refutando a necessidade obrigatória de clusters de alta "
    "performance para a prototipagem de modelos de linguagem específicos."
)

# Novo subtópico de reprodutibilidade — inserido após "Otimização e Regularização"
REPRO_TITLE = "Reprodutibilidade do Experimento"
REPRO_PARA = (
    "Para permitir a reprodução do experimento descrito, registram-se os seguintes "
    "parâmetros operacionais. Os checkpoints V6 e V7 foram gerados com o mesmo "
    "conjunto de hiperparâmetros reportado na Tabela 1, utilizando a biblioteca "
    "PyTorch (versão 2.11) sobre CUDA. O treinamento usou amostragem aleatória de "
    "batches (torch.randint com seed padrão do PyTorch) e gradient clipping com "
    "norma máxima 1,0. O agendamento da taxa de aprendizado seguiu o esquema "
    "CosineAnnealingLR, decaindo de 3e-4 para 1e-5 ao longo do total de iterações. "
    "O critério de parada adotado foi a estabilização da perda de validação, "
    "registrando-se o V6 ao final da iteração 1.400 e o V7 ao final da iteração 5.000, "
    "com salvamento contínuo de checkpoints a cada 200 iterações (parâmetro "
    "eval_iters). Reconhece-se como limitação o detalhamento parcial do pipeline — "
    "a fixação explícita de seed e a documentação granular do ambiente de execução "
    "(versões de biblioteca, variáveis ROCm/CUDA) ficam como recomendação para "
    "trabalhos subsequentes."
)

FONTE_TABELA = "Fonte: Resultados originais da pesquisa"
FONTE_FIGURA = "Fonte: Resultados originais da pesquisa"


# ---------------------------------------------------------------
# Main
# ---------------------------------------------------------------
def main():
    doc = Document(str(SRC))
    print(f"[carregado] {len(doc.paragraphs)} parágrafos, {len(doc.tables)} tabelas")

    # 1. Atualiza parágrafos com conteúdo desatualizado
    CONTENT_EDITS = {
        61: NOVO_PAR_OBJETIVOS,
        103: NOVO_PAR_103,
    }
    for idx, new_text in CONTENT_EDITS.items():
        set_paragraph_text(doc.paragraphs[idx], new_text)
        print(f"[ok] parágrafo {idx} atualizado")

    # 2. Corrige texto em outros parágrafos (30.58 -> 52,93)
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        new = t
        new = new.replace("30,58 milhões", "52,93 milhões")
        new = new.replace("30.58 milhões", "52,93 milhões")
        new = new.replace("30,58 Milhões", "52,93 Milhões")
        new = new.replace("30.58 Milhões", "52,93 Milhões")
        if new != t:
            set_paragraph_text(p, new)
            print(f"[ok] correção 30,58 -> 52,93 no parágrafo {i}")

    # 3. Encontra índice onde começam as Referências (não trocar siglas lá)
    refs_start = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "REFERÊNCIAS":
            refs_start = i
            break
    print(f"[info] Referências começam no parágrafo {refs_start}")

    # 4. Troca siglas (XXX) por [XXX] na primeira ocorrência
    replace_acronyms_first_occurrence(doc, refs_start)

    # 5. Localiza o parágrafo "SUMÁRIO" e remove ele + os 7 parágrafos seguintes (as entradas)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "SUMÁRIO":
            # remove o parágrafo do sumário e os próximos 7 (CONSIDERAÇÕES, IMPLEMENTAÇÃO, etc.)
            # NÃO removemos — apenas esvaziamos, pra não quebrar referências internas
            # Na verdade, esses "parágrafos seguintes" são o início do corpo do texto (CONSIDERAÇÕES INICIAIS).
            # Como o corpo inicia ali, vamos APENAS remover o parágrafo "SUMÁRIO" em si.
            parent = p._element.getparent()
            parent.remove(p._element)
            print(f"[ok] parágrafo SUMÁRIO removido (índice {i})")
            break

    # 6. Inserir subtópico de Reprodutibilidade depois de "Otimização e Regularização"
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Otimização e Regularização":
            # Queremos inserir APÓS o parágrafo seguinte (que é o texto da subseção)
            # Encontra o próximo parágrafo com estilo de título ou vazio
            insert_after = i + 1  # default: após o título
            # Procura fim do parágrafo da subseção (próximo heading ou um que seja RESULTADOS)
            for j in range(i + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip() == "RESULTADOS E DISCUSSÃO":
                    insert_after = j  # inserimos ANTES de "RESULTADOS"
                    break
            # Insere novo título + conteúdo antes do insert_after
            target_para = doc.paragraphs[insert_after]
            new_title = target_para.insert_paragraph_before(REPRO_TITLE)
            for r in new_title.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(11)
            new_body = target_para.insert_paragraph_before(REPRO_PARA)
            print(f"[ok] subtópico 'Reprodutibilidade do Experimento' inserido antes de RESULTADOS (índice {insert_after})")
            break

    # 7. Ajusta caption da Figura 1: título + fonte.
    #    O parágrafo 95 originalmente era a legenda. Agora precisamos também da Fonte logo abaixo.
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Figura 1.") and "Evolução" in p.text:
            # Adiciona uma linha imediatamente após com "Fonte: ..."
            # Verifica se já existe Fonte logo abaixo
            next_p = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if next_p and "Fonte:" not in next_p.text:
                new_fonte = p._element.addnext(OxmlElement("w:p"))
                # Não dá pra facilmente criar um paragraph via python-docx após addnext.
                # Em vez disso, inserimos antes do próximo.
                if next_p:
                    fonte_para = next_p.insert_paragraph_before(FONTE_FIGURA)
                    for r in fonte_para.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(11)
                    print(f"[ok] Fonte: adicionada abaixo da Figura 1")
            break

    # 8. Adiciona "Fonte: Resultados originais da pesquisa" abaixo das tabelas (se já não tiver)
    # No doc já existe "Fonte: Elaborada pelo autor (2025)" em [85] e [99].
    for i, p in enumerate(doc.paragraphs):
        if "Elaborada pelo autor" in p.text:
            set_paragraph_text(p, FONTE_TABELA)
            print(f"[ok] Fonte da tabela normalizada em [{i}]")

    # 9. Ajusta formatação global (margens, fonte, espaçamento)
    apply_global_formatting(doc)
    print("[ok] formatação global aplicada (margens 2,5cm, Arial 11, espaçamento 1,5)")

    # 10. Adiciona paginação no rodapé
    add_footer_page_number(doc)
    print("[ok] paginação Arial 9 no canto inferior direito")

    # Salva
    doc.save(str(DST))
    print(f"\n=== SALVO: {DST} ===")
    print("No Word/LibreOffice:")
    print("  1. Abra o arquivo e pressione Ctrl+A → F9 para atualizar campos (paginação)")
    print("  2. Verifique cabeçalho (já deve vir preservado do original)")
    print("  3. Revise a nova subseção 'Reprodutibilidade do Experimento'")
    print("  4. Revise o novo parágrafo de objetivos (61)")
    print("  5. Confira que as tabelas estão sem preenchimento colorido/negrito")


if __name__ == "__main__":
    main()
