"""Aplica os 6 pontos de correção da orientadora + atualização de hardware real.

Pontos:
  1. Mistura Metodologia/Resultados: suavizar interpretações em Methodology
     e concentrá-las em Discussão (dentro de Resultados)
  2. Adicionar parágrafo na Conclusão sobre tarefas downstream não avaliadas
  3. Tabela 2: tornar explícita a separação V6 vs V7; reescrever texto de
     Resultados explicando a "regressão" aparente da loss
  4. Reprodutibilidade: atualizar hardware (NVIDIA RTX 3050), split 90/10,
     menção explícita de seed não fixa
  5. Trocar "early stopping implícita" por "parada baseada na estabilização
     empírica da perda de validação"
  6. Adicionar Abstract em inglês logo após as Palavras-chave
"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path("/tmp/tcc_v3.docx")
DST = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (template).docx")


# ---------------------------------------------------------------
# Textos novos
# ---------------------------------------------------------------
NEW_53 = (
    "O treinamento foi executado em uma estação de trabalho local equipada com GPU "
    "NVIDIA GeForce RTX 3050 Laptop (6 GB de VRAM), CPU Intel Core i5-13420H e 7 GB de "
    "memória RAM. A capacidade limitada de memória de vídeo foi o fator determinante "
    "para o dimensionamento da arquitetura, restringindo o tamanho do lote (batch size) "
    "e a profundidade da rede. O código foi implementado em Python 3.12.3 utilizando a "
    "biblioteca PyTorch 2.11 com suporte à aceleração via CUDA."
)

NEW_61 = (
    "O dataset final foi tokenizado utilizando o vocabulário BPE (Byte-Pair Encoding) "
    "padrão do GPT-2, de 50.257 tokens, resultando em um volume total de 2.506.590 "
    "tokens de treino. O vocabulário do GPT-2 foi reutilizado em vez de ser treinado um "
    "tokenizador específico para o português, escolha motivada pela simplificação do "
    "pipeline e pela compatibilidade com ferramentas pré-existentes. O retreinamento de "
    "um tokenizador próprio em corpus PT-BR é indicado como trabalho futuro."
)

NEW_91 = (
    "Para permitir a reprodução do experimento descrito, registram-se os seguintes "
    "parâmetros operacionais. Os checkpoints V6 e V7 foram gerados com o conjunto de "
    "hiperparâmetros reportado na Tabela 1, utilizando a biblioteca PyTorch 2.11 sobre "
    "CUDA em GPU NVIDIA GeForce RTX 3050 Laptop (6 GB de VRAM), CPU Intel Core "
    "i5-13420H e 7 GB de memória RAM. O split entre treino e validação seguiu proporção "
    "de 90%/10%. O experimento não utilizou seed fixa para a amostragem aleatória de "
    "batches nem para a inicialização dos pesos, o que pode introduzir variação nos "
    "resultados entre execuções; recomenda-se, em replicações futuras, a fixação "
    "explícita de seed para maior determinismo. O treinamento empregou amostragem "
    "aleatória de batches e gradient clipping com norma máxima 1,0. O agendamento da "
    "taxa de aprendizado seguiu o esquema CosineAnnealingLR, decaindo de 3e-4 para 1e-5 "
    "ao longo do total de iterações. O critério de parada adotado foi a parada baseada "
    "na estabilização empírica da perda de validação, registrando-se o V6 ao final da "
    "iteração 1.400 e o V7 ao final da iteração 5.000, com salvamento contínuo de "
    "checkpoints a cada 200 iterações."
)

NEW_103 = (
    "A Tabela 2 apresenta as métricas de dois experimentos distintos: V6, treinado "
    "sobre o corpus inicial composto apenas por textos literários e normativos "
    "(aproximadamente 6 MB de texto), e V7, treinado sobre o corpus expandido que "
    "inclui pares pergunta/resposta em português (aproximadamente 37 MB de texto). Para "
    "o V6, a perda de validação iniciou em 4,286 na iteração 200 e apresentou "
    "decaimento consistente até atingir 2,145 na iteração 1.400, com acurácia de "
    "predição de token de 63,5%. Para o V7, após treinamento sobre o corpus expandido, "
    "a perda de validação estabilizou em 2,003 na iteração 5.000, com acurácia de "
    "59,4%."
)

NEW_103_DISC = (
    "Esse comportamento pode ser explicado por duas dinâmicas distintas. No V6, o "
    "tamanho reduzido do corpus permitiu que o modelo memorizasse porções "
    "significativas do conjunto de treino, resultando em perda de treino de 1,044 na "
    "iteração 1.400 e gap train/val próximo de 1,10 unidades, sinal clássico de "
    "sobreajuste (overfitting). No V7, a ampliação do corpus introduziu maior "
    "heterogeneidade e inibiu a memorização, elevando a perda de treino para a faixa "
    "de 1,78 a 1,84 entre as iterações 3.000 e 5.000, porém com gap train/val reduzido "
    "a 0,22 unidades, comportamento consistente com generalização. A aparente "
    "regressão da perda de treino entre V6 (1,044) e V7 (1,835) reflete a mudança de "
    "regime entre os dois experimentos e não uma degradação do aprendizado."
)

NEW_103_TOK = (
    "Cabe ainda discutir o impacto da escolha do tokenizador GPT-2. Esse vocabulário, "
    "treinado majoritariamente em inglês, segmenta palavras comuns do português como "
    "\"não\", \"são\" e \"coração\" em múltiplos subtokens, aumentando o comprimento "
    "efetivo das sequências e reduzindo a quantidade de conteúdo útil comportado na "
    "janela de contexto de 256 tokens. Esse efeito contribui parcialmente para a "
    "limitação observada de coerência semântica em gerações longas."
)

NEW_DOWNSTREAM = (
    "Embora o modelo tenha sido concebido com foco em aplicações como atendimento ao "
    "cliente e resumo de documentos, esta etapa do trabalho concentrou-se "
    "exclusivamente no pré-treino e avaliação intrínseca, por meio das métricas de "
    "perda de validação e acurácia de predição de token. A avaliação em tarefas "
    "downstream específicas, incluindo métricas automáticas como ROUGE-L e BERTScore "
    "para a tarefa de resumo de documentos e avaliação humana de respostas geradas em "
    "cenário de atendimento ao cliente, permanece como etapa futura, viabilizada pelo "
    "pipeline de treinamento e pela arquitetura entregues."
)

ABSTRACT_TEXT = (
    "This work designs, implements, and trains, from scratch, a Small Language Model "
    "(SLM) decoder-only in Brazilian Portuguese, executed on academic hardware with 6 "
    "GB of VRAM. The method comprises: curation of a multi-domain corpus (legal, "
    "classical literature, and contemporary fiction); use of the GPT-2 BPE tokenizer as "
    "baseline (50,257 tokens); implementation of a Transformer decoder-only "
    "architecture with 52.93 million parameters (8 layers, 6 attention heads, "
    "384-dimensional embeddings, 256-token context window); causal pre-training with "
    "AdamW optimizer, cosine learning rate schedule, and checkpointing based on "
    "empirical validation-loss stabilization; and quantitative evaluation through "
    "validation loss and next-token prediction accuracy, complemented by qualitative "
    "evaluation of generated samples. Results show progressive convergence, with "
    "validation loss reduced from 4.29 (step 200) to 2.00 (step 5000), with "
    "corresponding accuracy gain from 26.2% to 59.4%. The model learns syntactic "
    "structure, verbal conjugation, and dialogue format of Portuguese, although it "
    "maintains limited semantic coherence in long sequences, consistent with the "
    "dataset scale (2.5 million tokens) and available parameters. As a direct "
    "follow-up, the work incorporates instruction datasets translated to PT-BR (Dolly "
    "15k and Alpaca) into the corpus to enable, in a subsequent version, supervised "
    "fine-tuning directed at customer service and document summarization. Limitations "
    "include corpus size, the use of a tokenizer not retrained for Portuguese, and the "
    "absence of preference-based alignment."
)
KEYWORDS_EN = "Keywords: Portuguese; Pre-training; Transformer; SLM; BPE Tokenization."


# ---------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------
def set_paragraph_simple_text(elem, new_text: str):
    pPr = elem.find(qn("w:pPr"))
    first_run = None
    runs_to_remove = []
    for r in elem.iter(qn("w:r")):
        if first_run is None:
            first_run = r
        else:
            runs_to_remove.append(r)
    for r in runs_to_remove:
        r.getparent().remove(r)
    if first_run is not None:
        ts = list(first_run.iter(qn("w:t")))
        for t in ts[1:]:
            t.getparent().remove(t)
        if ts:
            ts[0].text = new_text
        else:
            t = OxmlElement("w:t")
            t.text = new_text
            first_run.append(t)


def set_paragraph_text_as_new(paragraph, text: str):
    """Substitui texto preservando estilo via docx API."""
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_after_with_style(ref_paragraph, text: str, style=None, bold=False):
    """Insere parágrafo APÓS o referenciado. Retorna o novo."""
    # python-docx tem apenas insert_paragraph_before. Usa addnext via lxml.
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    # Pega como Paragraph real
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, ref_paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    run = new_para.add_run(text)
    if bold:
        run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return new_para


# ---------------------------------------------------------------
# Main
# ---------------------------------------------------------------
def main():
    doc = Document(str(SRC))
    print(f"[carregado] {len(doc.paragraphs)} parágrafos\n")

    # Localiza parágrafos por texto (robusto a mudanças de índice)
    def find_idx_starting_with(prefix, after_idx=0):
        for i in range(after_idx, len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip().startswith(prefix):
                return i
        return None

    def find_idx_containing(substr, after_idx=0):
        for i in range(after_idx, len(doc.paragraphs)):
            if substr.lower() in doc.paragraphs[i].text.lower():
                return i
        return None

    # 1. Recursos Computacionais (AMD → NVIDIA)
    i_hw = find_idx_containing("AMD Radeon RX 5600")
    if i_hw is not None:
        set_paragraph_text_as_new(doc.paragraphs[i_hw], NEW_53)
        doc.paragraphs[i_hw].paragraph_format.first_line_indent = Cm(1.25)
        doc.paragraphs[i_hw].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        print(f"[ok 1] par {i_hw} (Recursos Computacionais) atualizado — GPU NVIDIA")

    # 2. Tokenização — remove interpretação ("introduz uma limitação")
    i_tok = find_idx_containing("introduz uma limitação")
    if i_tok is not None:
        set_paragraph_text_as_new(doc.paragraphs[i_tok], NEW_61)
        doc.paragraphs[i_tok].paragraph_format.first_line_indent = Cm(1.25)
        doc.paragraphs[i_tok].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        print(f"[ok 2] par {i_tok} (Tokenização) reescrita sem interpretação")

    # 3. Reprodutibilidade — hardware + seed + split
    i_rep = find_idx_containing("reprodução do experimento")
    if i_rep is not None:
        set_paragraph_text_as_new(doc.paragraphs[i_rep], NEW_91)
        doc.paragraphs[i_rep].paragraph_format.first_line_indent = Cm(1.25)
        doc.paragraphs[i_rep].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        print(f"[ok 3] par {i_rep} (Reprodutibilidade) atualizado")

    # 4. Resultados — reescrever com separação V6/V7 + discussão
    i_res = find_idx_containing("perda de validação iniciou em 4,286")
    if i_res is not None:
        set_paragraph_text_as_new(doc.paragraphs[i_res], NEW_103)
        doc.paragraphs[i_res].paragraph_format.first_line_indent = Cm(1.25)
        doc.paragraphs[i_res].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Insere o parágrafo de discussão logo depois
        p_disc = insert_paragraph_after_with_style(doc.paragraphs[i_res], NEW_103_DISC)
        p_disc.paragraph_format.first_line_indent = Cm(1.25)
        p_disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Discussão sobre tokenizer (tirada da Metodologia)
        p_tok_disc = insert_paragraph_after_with_style(p_disc, NEW_103_TOK)
        p_tok_disc.paragraph_format.first_line_indent = Cm(1.25)
        p_tok_disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        print(f"[ok 4] par {i_res} (Resultados) reescrito + 2 parágrafos de discussão adicionados")

    # 5. Parágrafo da Conclusão — adicionar frase sobre downstream tasks
    # Queremos adicionar ao FINAL da Conclusão (antes de Referências)
    i_refs = find_idx_containing("Referências")
    # Pega o parágrafo imediatamente antes de Referências
    if i_refs is not None and i_refs > 0:
        # Insere um parágrafo ANTES de "Referências" com o texto novo
        ref_p = doc.paragraphs[i_refs]
        new_p = ref_p.insert_paragraph_before(NEW_DOWNSTREAM)
        new_p.paragraph_format.first_line_indent = Cm(1.25)
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in new_p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(11)
        print(f"[ok 5] parágrafo de tarefas downstream inserido antes de Referências")

    # 6. Abstract em inglês — inserir depois de Palavras-chave
    i_kw = find_idx_starting_with("Palavras-chave")
    if i_kw is not None:
        # Título "Abstract" em negrito
        p_abs_title = insert_paragraph_after_with_style(doc.paragraphs[i_kw], "Abstract", bold=True)
        # Corpo do Abstract
        p_abs_body = insert_paragraph_after_with_style(p_abs_title, ABSTRACT_TEXT)
        # Keywords
        p_kw_en = insert_paragraph_after_with_style(p_abs_body, KEYWORDS_EN, bold=False)
        # Negrito só em "Keywords:"
        p_kw_en.runs[0].text = ""
        r_bold = p_kw_en.add_run("Keywords: ")
        r_bold.bold = True
        r_bold.font.name = "Arial"; r_bold.font.size = Pt(11)
        r_rest = p_kw_en.add_run("Portuguese; Pre-training; Transformer; SLM; BPE Tokenization.")
        r_rest.font.name = "Arial"; r_rest.font.size = Pt(11)
        print(f"[ok 6] Abstract em inglês + Keywords inseridos após Palavras-chave")

    # 7. Tabela 2 — tornar a separação V6/V7 explícita nas linhas iniciais
    # Atual: linha 1 "200" (V6 implícito), linha 5 "1.400 (V6)"
    # Novo: "200 (V6)", "400 (V6)", "600 (V6)", "800 (V6)"
    t2 = doc.tables[1]
    if len(t2.rows) >= 5 and len(t2.rows[1].cells) >= 1:
        # Renomeia as primeiras 4 linhas de dados adicionando (V6)
        for row_idx, step_label in [(1, "200 (V6)"), (2, "400 (V6)"),
                                     (3, "600 (V6)"), (4, "800 (V6)")]:
            cell = t2.rows[row_idx].cells[0]
            # Limpa e reescreve preservando formatação
            for p in cell.paragraphs[1:]:
                p._element.getparent().remove(p._element)
            p0 = cell.paragraphs[0]
            set_paragraph_text_as_new(p0, step_label)
        print(f"[ok 7] Tabela 2: linhas 200-800 marcadas como (V6)")

    # Salva
    doc.save(str(DST))
    print(f"\n=== SALVO: {DST} ===")


if __name__ == "__main__":
    main()
