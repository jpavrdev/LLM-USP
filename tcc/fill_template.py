"""Preenche o template oficial USP/Esalq com o conteúdo real do TCC.

Estratégia: o template oficial é a base (preserva TODOS os estilos, margens,
cabeçalhos, paginação). Substituímos apenas os placeholders pelo conteúdo real
do `TCC MBA - USP 2025 (final).docx`, mantendo a formatação do template.

Input:
  - Template: /home/akirem/Downloads/Template TCC - Implementação de Algoritmo(s) de Machine Learning.docx
  - Fonte conteúdo: /home/akirem/Downloads/TCC MBA - USP 2025 (final).docx
Output:
  - /home/akirem/Downloads/TCC MBA - USP 2025 (template).docx
"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


TEMPLATE = Path("/home/akirem/Downloads/Template TCC - Implementação de Algoritmo(s) de Machine Learning.docx")
SOURCE = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (final).docx")
OUTPUT = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (template).docx")


TITLE_WORK = "Geração automática de texto com LLMs para atendimento ao cliente ou resumo de documentos"
AUTHOR = "João Pedro Azevedo Veras dos Reis"
ORIENTADOR = "Gabrielle Lombardi"
EMAIL_AUTOR = "pedro.veras@usp.br"

# Endereços — mínimo exigido pelo template: titulação/função + endereço/cidade/estado/país
ADDR_ALUNO = ("Universidade de São Paulo, Escola Superior de Agricultura Luiz de Queiroz. "
              "Discente do MBA em Data Science and Analytics. "
              "Av. Pádua Dias, 11 - Caixa Postal 9, Agência Centro, Piracicaba, SP, 13418-900, Brasil")
ADDR_ORIENT = ("Universidade de São Paulo, Escola Superior de Agricultura Luiz de Queiroz. "
               "Orientadora do MBA. "
               "Departamento de Economia, Administração e Sociologia, Av. Pádua Dias, 11, "
               "Piracicaba, SP, 13418-900, Brasil")

# Títulos canônicos das seções (idênticos ao template)
SEC_TITLES_TEMPLATE = [
    "Considerações Iniciais",
    "Implementação de Algoritmo(s) de Machine Learning",
    "Resultados e Discussão",
    "Conclusão(ões)",
    "Agradecimento (opcional, 1 parágrafo, bem sucinto)",
    "Referências",
    "Apêndice ou Anexo (opcional)",
]

# Mapeamento: título no source → título no template (quando diferem)
SOURCE_TO_TEMPLATE_TITLE = {
    "Considerações Iniciais": "Considerações Iniciais",
    "Implementação de Algoritmo(s) de Machine Learning": "Implementação de Algoritmo(s) de Machine Learning",
    "Resultados e Discussão": "Resultados e Discussão",
    "Conclusão": "Conclusão(ões)",
    "Referências": "Referências",
}


# ---------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------
def get_text(elem) -> str:
    """Texto bruto de um elemento <w:p>."""
    if not elem.tag.endswith("}p"):
        return ""
    return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()


def body_children(doc):
    """Filhos diretos do body (parágrafos + tabelas em ordem)."""
    return list(doc.element.body)


def find_para_index_by_text(doc, needle_text, prefix_match=False) -> int | None:
    """Retorna o índice (em body children) do primeiro parágrafo cujo texto bate."""
    for i, child in enumerate(body_children(doc)):
        if not child.tag.endswith("}p"):
            continue
        t = get_text(child)
        if prefix_match:
            if t.startswith(needle_text):
                return i
        else:
            if t == needle_text:
                return i
    return None


def set_paragraph_simple_text(elem, new_text: str):
    """Substitui todo o texto de um <w:p>, preservando o pPr e a primeira run como template de formatação."""
    # Remove todas as runs (<w:r>) mas mantém pPr
    pPr = elem.find(qn("w:pPr"))
    first_run = None
    runs_to_remove = []
    for r in elem.iter(qn("w:r")):
        if first_run is None:
            first_run = r
        else:
            runs_to_remove.append(r)
    for r in runs_to_remove:
        r.getparent().remove(r)
    if first_run is not None:
        # Limpa os <w:t> da primeira run, deixa só um com o novo texto
        ts = list(first_run.iter(qn("w:t")))
        for t in ts[1:]:
            t.getparent().remove(t)
        if ts:
            ts[0].text = new_text
        else:
            # Cria novo <w:t>
            from docx.oxml import OxmlElement
            t = OxmlElement("w:t")
            t.text = new_text
            first_run.append(t)


def extract_section_from_source(src_doc, section_title, next_section_title) -> list:
    """Pega a lista de elementos XML do conteúdo de uma seção no TCC-fonte.
    Exclui o parágrafo do título em si."""
    elements = []
    started = False
    for child in body_children(src_doc):
        if child.tag.endswith("}p"):
            t = get_text(child)
            if t == section_title:
                started = True
                continue
            if t == next_section_title:
                break
        if started:
            elements.append(deepcopy(child))
    return elements


def replace_section_body(template_doc, template_title, content_elements):
    """No template, remove tudo entre o título da seção e o próximo título
    (removendo placeholders/atenções) e injeta os content_elements."""
    body = template_doc.element.body
    children = body_children(template_doc)
    i_title = None
    for i, child in enumerate(children):
        if get_text(child) == template_title:
            i_title = i
            break
    if i_title is None:
        print(f"[warn] título {template_title!r} não encontrado no template")
        return
    # Encontra fim da seção (próximo título da lista oficial)
    i_next = None
    for j in range(i_title + 1, len(children)):
        if get_text(children[j]) in SEC_TITLES_TEMPLATE:
            i_next = j
            break
    if i_next is None:
        i_next = len(children)
    # Remove TODOS os elementos entre o título e o próximo título
    for elem in children[i_title + 1 : i_next]:
        body.remove(elem)
    # Insere os novos elementos após o título
    title_elem = body_children(template_doc)[i_title]
    prev = title_elem
    for new_elem in content_elements:
        prev.addnext(new_elem)
        prev = new_elem
    print(f"[ok] seção '{template_title}' preenchida com {len(content_elements)} elementos")


def remove_section(template_doc, section_title):
    """Remove uma seção inteira (título + conteúdo até o próximo título)."""
    body = template_doc.element.body
    children = body_children(template_doc)
    i_start = None
    for i, child in enumerate(children):
        if get_text(child) == section_title:
            i_start = i
            break
    if i_start is None:
        return
    i_end = None
    for j in range(i_start + 1, len(children)):
        if get_text(children[j]) in SEC_TITLES_TEMPLATE:
            i_end = j
            break
    if i_end is None:
        i_end = len(children)
    for elem in children[i_start:i_end]:
        body.remove(elem)
    print(f"[ok] seção '{section_title}' removida")


# ---------------------------------------------------------------
# Main
# ---------------------------------------------------------------
def main():
    template = Document(str(TEMPLATE))
    source = Document(str(SOURCE))
    print(f"[template] {len(template.paragraphs)} parágrafos")
    print(f"[source]   {len(source.paragraphs)} parágrafos\n")

    # ---------------------------------------------------------------
    # FOLHA DE ROSTO — substituir placeholders por dados reais
    # ---------------------------------------------------------------
    # Par 0: Título do trabalho (primeiro) — começa com "Título do trabalho de conclusão..."
    for p in template.paragraphs:
        if p.text.strip().startswith("Título do trabalho"):
            set_paragraph_simple_text(p._element, TITLE_WORK)
            print(f"[ok] título 1 substituído")
            break
    # Par ~28: Título do trabalho (segundo, antes do Resumo)
    count = 0
    for p in template.paragraphs:
        if p.text.strip().startswith("Título do trabalho"):
            count += 1
            if count == 2:
                set_paragraph_simple_text(p._element, TITLE_WORK)
                print(f"[ok] título 2 substituído (antes do Resumo)")
                break

    # Autores
    for p in template.paragraphs:
        t = p.text.strip()
        if "nome completo aluno" in t.lower() and "orientador" in t.lower():
            set_paragraph_simple_text(p._element, f"{AUTHOR}¹*; {ORIENTADOR}²")
            print(f"[ok] autores substituídos")
            break

    # Endereços (2 parágrafos: 1 Nome da Empresa... e 2 Nome da Empresa...)
    addr_replacements = [
        (lambda t: t.startswith("1 Nome da Empresa"), f"¹ {ADDR_ALUNO}"),
        (lambda t: t.startswith("2 Nome da Empresa"), f"² {ADDR_ORIENT}"),
    ]
    for cond, replacement in addr_replacements:
        for p in template.paragraphs:
            if cond(p.text.strip()):
                set_paragraph_simple_text(p._element, replacement)
                print(f"[ok] endereço substituído: {replacement[:50]}...")
                break

    # Email do autor
    for p in template.paragraphs:
        if p.text.strip().startswith("*autor correspondente"):
            set_paragraph_simple_text(p._element, f"*autor correspondente: {EMAIL_AUTOR}")
            print(f"[ok] email substituído")
            break

    # ---------------------------------------------------------------
    # RESUMO — substituir placeholder
    # ---------------------------------------------------------------
    # O placeholder começa com "O resumo é uma descrição geral..."
    # O resumo real está no source (primeiro parágrafo após "Resumo")
    src_resumo_idx = None
    for i, p in enumerate(source.paragraphs):
        if p.text.strip() == "Resumo":
            src_resumo_idx = i + 1
            break
    # No source, após "Resumo", vem direto o texto (1 parágrafo) + palavras-chave
    if src_resumo_idx is not None:
        resumo_texto = source.paragraphs[src_resumo_idx].text
        # Palavras-chave está no próximo parágrafo
        palavras_chave_texto = None
        for j in range(src_resumo_idx + 1, min(src_resumo_idx + 5, len(source.paragraphs))):
            t = source.paragraphs[j].text.strip()
            if t.upper().startswith("PALAVRAS-CHAVE") or t.startswith("Palavras-chave"):
                palavras_chave_texto = t
                break

        # Substitui placeholder do resumo
        for p in template.paragraphs:
            if p.text.strip().startswith("O resumo é uma descrição"):
                set_paragraph_simple_text(p._element, resumo_texto)
                print(f"[ok] resumo preenchido")
                break

        # Palavras-chave
        if palavras_chave_texto:
            # Padroniza: "Palavras-chave: X; Y; Z."
            if palavras_chave_texto.upper().startswith("PALAVRAS-CHAVE"):
                # Já está no formato, só troca Palavras-chave
                palavras_chave_texto = palavras_chave_texto.replace("PALAVRAS-CHAVE", "Palavras-chave", 1)
            for p in template.paragraphs:
                if p.text.strip().startswith("Palavras-chave:"):
                    set_paragraph_simple_text(p._element, palavras_chave_texto)
                    print(f"[ok] palavras-chave preenchidas")
                    break

    # ---------------------------------------------------------------
    # REMOVER AVISOS "Atenção: antes de enviar o arquivo..."
    # ---------------------------------------------------------------
    to_remove = []
    for p in template.paragraphs:
        if p.text.strip().startswith("Atenção: antes de enviar"):
            to_remove.append(p._element)
    for elem in to_remove:
        elem.getparent().remove(elem)
    print(f"[ok] {len(to_remove)} avisos 'Atenção:' removidos")

    # ---------------------------------------------------------------
    # REMOVER bloco opcional Abstract/Keywords se presente
    # ---------------------------------------------------------------
    # Remove desde "Título em inglês ou espanhol (opcional)" até antes de "Considerações Iniciais"
    body = template.element.body
    children = body_children(template)
    i_abstract = None
    i_consid = None
    for i, c in enumerate(children):
        t = get_text(c)
        if t.startswith("Título em inglês ou espanhol") and i_abstract is None:
            i_abstract = i
        if t == "Considerações Iniciais":
            i_consid = i
            break
    if i_abstract is not None and i_consid is not None and i_abstract < i_consid:
        for elem in children[i_abstract:i_consid]:
            body.remove(elem)
        print(f"[ok] bloco Abstract/Resumen/Keywords opcional removido ({i_consid - i_abstract} elementos)")

    # ---------------------------------------------------------------
    # SEÇÕES PRINCIPAIS — substituir placeholder pelo conteúdo real
    # ---------------------------------------------------------------
    for src_title, tpl_title in SOURCE_TO_TEMPLATE_TITLE.items():
        # Determina o próximo título no SOURCE
        all_src_titles = list(SOURCE_TO_TEMPLATE_TITLE.keys())
        idx = all_src_titles.index(src_title)
        next_src_title = all_src_titles[idx + 1] if idx + 1 < len(all_src_titles) else "___FIM___"
        # Extrai elementos
        content = extract_section_from_source(source, src_title, next_src_title)
        # Substitui no template
        replace_section_body(template, tpl_title, content)

    # ---------------------------------------------------------------
    # REMOVER seções opcionais (Agradecimento e Apêndice)
    # ---------------------------------------------------------------
    remove_section(template, "Agradecimento (opcional, 1 parágrafo, bem sucinto)")
    remove_section(template, "Apêndice ou Anexo (opcional)")

    # Salva
    template.save(str(OUTPUT))
    print(f"\n=== SALVO: {OUTPUT} ===")
    print(f"Total de parágrafos: {len(template.paragraphs)}")
    print(f"Total de tabelas: {len(template.tables)}")


if __name__ == "__main__":
    main()
