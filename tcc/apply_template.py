"""Aplica o template oficial USP/Esalq no TCC:
  - Títulos de seção de CAIXA ALTA → Title Case conforme template
  - Estilo 'List Paragraph' nos títulos de seção
  - Inserção do título do trabalho antes do RESUMO (página 2)
  - Preserva: margens, paginação, recuos do corpo, page breaks
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


SRC = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (final).docx")
DST = SRC

TITLE_WORK = "Geração automática de texto com LLMs para atendimento ao cliente ou resumo de documentos"

# Mapeamento de títulos: CAIXA ALTA → Title Case
TITLE_MAP = {
    "RESUMO": "Resumo",
    "CONSIDERAÇÕES INICIAIS": "Considerações Iniciais",
    "IMPLEMENTAÇÃO DE ALGORITMO(S) DE MACHINE LEARNING": "Implementação de Algoritmo(s) de Machine Learning",
    "RESULTADOS E DISCUSSÃO": "Resultados e Discussão",
    "CONCLUSÃO": "Conclusão",
    "REFERÊNCIAS": "Referências",
}


def set_paragraph_text(paragraph, text: str):
    """Troca texto preservando as runs (aplica ao primeiro run)."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def main():
    doc = Document(str(SRC))
    print(f"[carregado] {len(doc.paragraphs)} parágrafos\n")

    # 1) Renomear títulos de seção e aplicar estilo "List Paragraph"
    try:
        list_para_style = doc.styles["List Paragraph"]
    except KeyError:
        list_para_style = None
        print("[aviso] estilo 'List Paragraph' não existe neste doc, criando via add_style")

    for p in doc.paragraphs:
        original = p.text.strip()
        if original in TITLE_MAP:
            new_text = TITLE_MAP[original]
            # Preserva page_break_before
            had_page_break = p.paragraph_format.page_break_before
            set_paragraph_text(p, new_text)
            # Aplica negrito em todos os runs
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(11)
            # Alinhamento à esquerda
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Tenta aplicar estilo List Paragraph
            if list_para_style is not None:
                try:
                    p.style = list_para_style
                    # Após mudança de estilo, reaplica negrito/alinhamento
                    for r in p.runs:
                        r.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except Exception as e:
                    print(f"[aviso] não pude aplicar List Paragraph em {new_text!r}: {e}")
            # Restaura page break before
            if had_page_break:
                p.paragraph_format.page_break_before = True
            # Remove numeração/lista se estilo List Paragraph adicionou automaticamente
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    pPr.remove(numPr)
            # Remove indentação (List Paragraph vem com indent)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.first_line_indent = Cm(0)
            print(f"[ok] {original!r} → {new_text!r}")

    # 2) Inserir título do trabalho antes do RESUMO (se ainda não estiver lá)
    # Procura o parágrafo "Resumo" e verifica se já há um título igual ao do trabalho logo antes.
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Resumo":
            # Olha os 3 parágrafos anteriores pra ver se já tem o título
            tem_titulo_repetido = False
            for j in range(max(0, i - 5), i):
                if TITLE_WORK.lower() in doc.paragraphs[j].text.lower():
                    tem_titulo_repetido = True
                    break
            if not tem_titulo_repetido:
                # Insere o título antes do parágrafo "Resumo"
                new_p = p.insert_paragraph_before(TITLE_WORK)
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in new_p.runs:
                    r.bold = True
                    r.font.name = "Arial"
                    r.font.size = Pt(11)
                # Garante que esse parágrafo herde o page break antes do Resumo
                # (e remove o page break do Resumo pra não criar duas páginas)
                if p.paragraph_format.page_break_before:
                    new_p.paragraph_format.page_break_before = True
                    p.paragraph_format.page_break_before = False
                print(f"[ok] título do trabalho inserido antes do Resumo")
            break

    doc.save(str(DST))
    print(f"\n=== SALVO: {DST} ===")

    # Validação
    doc2 = Document(str(DST))
    print("\n=== Títulos após correção ===")
    for i, p in enumerate(doc2.paragraphs):
        t = p.text.strip()
        if t in TITLE_MAP.values() or t == TITLE_WORK:
            pb = "[PG]" if p.paragraph_format.page_break_before else ""
            print(f"  [{i}] style={p.style.name!r} {pb} {t!r}")


if __name__ == "__main__":
    main()
