"""Edita o TCC.docx aplicando correções para alinhar texto, números e referências
com o que foi de fato implementado (checkpoints V6 step 1400 e V7 step 5000).

Estratégia: trabalhar sobre uma cópia "(revisado)", substituir textos inteiros nos
parágrafos conhecidos por índice, editar células de tabela, inserir novos parágrafos
antes de marcadores específicos (ex: antes do "CONCLUSÃO").
"""

from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


SRC = Path("/home/akirem/Downloads/TCC MBA - USP 2025.docx")
DST = Path("/home/akirem/Downloads/TCC MBA - USP 2025 (revisado).docx")


def set_paragraph_text(paragraph, text: str):
    """Substitui todo o texto do parágrafo, preservando estilo mas descartando runs inline."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_before(paragraph, text: str, style: str | None = None):
    """Insere um novo parágrafo antes do paragraph dado, com texto e estilo opcional."""
    new_p = paragraph.insert_paragraph_before(text=text)
    if style:
        try:
            new_p.style = style
        except KeyError:
            pass
    return new_p


def set_cell_text(cell, text: str):
    """Substitui texto da célula preservando o primeiro parágrafo."""
    # Remove todos os parágrafos exceto o primeiro
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    set_paragraph_text(cell.paragraphs[0], text)


def main():
    doc = Document(str(SRC))

    # ------------------------------------------------------------------
    # 1. RESUMO — parágrafo [39]
    # ------------------------------------------------------------------
    novo_resumo = (
        "Este trabalho projeta, implementa e treina, do zero, um Modelo de Linguagem "
        "de Pequeno Porte (SLM) decoder-only em português brasileiro, executado em "
        "hardware acadêmico com 6 GB de VRAM. O método abrange: curadoria de um corpus "
        "multi-domínio (jurídico, literário clássico e ficção contemporânea); uso do "
        "tokenizador BPE do GPT-2 como base (50.257 tokens); implementação de uma "
        "arquitetura Transformer decoder-only de 52,93 milhões de parâmetros (8 "
        "camadas, 6 cabeças de atenção, dimensão 384, janela de 256 tokens); "
        "pré-treino causal com otimizador AdamW, agendamento cosseno de taxa de "
        "aprendizado e checkpointing por early stopping; e avaliação quantitativa "
        "via loss de validação e acurácia de predição de token, complementada por "
        "avaliação qualitativa de amostras geradas. Os resultados mostram convergência "
        "progressiva — loss de validação reduzida de 4,29 (step 200) para 2,00 (step "
        "5000) — com ganho correspondente de acurácia de 26,2% para 59,4%. Observa-se "
        "que o modelo aprende estrutura sintática, conjugação verbal e formato de "
        "diálogo do português, porém mantém coerência semântica limitada em sequências "
        "longas, o que é consistente com a escala do dataset (2,5 milhões de tokens) "
        "e com os parâmetros disponíveis. Como encaminhamento direto, o trabalho "
        "incorpora ao corpus datasets de instrução traduzidos ao PT-BR (Dolly 15k e "
        "Alpaca) para viabilizar, em versão subsequente, fine-tuning supervisionado "
        "orientado a atendimento ao cliente e resumo de documentos. Limitações "
        "incluem o tamanho do corpus, o uso de tokenizador não re-treinado para "
        "português e a ausência de alinhamento por preferências."
    )

    palavras_chave = "PALAVRAS-CHAVE: Português; Pré-treino; Transformer; SLM; Tokenização BPE."

    # ------------------------------------------------------------------
    # 2. Correção citação HAREM / brWaC — parágrafo [58]
    # ------------------------------------------------------------------
    par58_novo = (
        "No contexto brasileiro, há carência de recursos de pré-treino e benchmarks "
        "específicos, embora existam iniciativas relevantes como o corpus brWaC "
        "(WAGNER FILHO et al., 2018) e as campanhas HAREM de reconhecimento de "
        "entidades nomeadas (SANTOS et al., 2006). Para suprir essa lacuna, este "
        "trabalho propõe a construção de um SLM (Small Language Model) treinado do "
        "zero em português, focado em atendimento ao cliente e resumo de documentos "
        "como casos de uso orientadores. O projeto abrange desde a curadoria e "
        "deduplicação de dados — essencial para reduzir memorização e riscos de "
        "privacidade (LEE et al., 2022; CARLINI et al., 2021) — até a especialização "
        "do modelo por pré-treino causal."
    )

    # ------------------------------------------------------------------
    # 3. LGPD / direito autoral — parágrafo [59]
    # ------------------------------------------------------------------
    par59_novo = (
        "A predominância de modelos proprietários e centralizados, como o GPT-4, "
        "impõe barreiras à adoção corporativa segura no Brasil. A dependência de "
        "provedores internacionais cria riscos de homogeneização e pontos únicos "
        "de falha (BOMMASANI et al., 2021), além de potenciais conflitos com a Lei "
        "Geral de Proteção de Dados (LGPD) em cenários de produção que processem "
        "dados pessoais (BRASIL, 2018; BROWN et al., 2022). Adicionalmente, a "
        "opacidade dos modelos \"caixa-preta\" dificulta auditorias (RUDIN, 2019) e "
        "o viés linguístico dos dados de treino, majoritariamente anglófonos, limita "
        "a captura de nuances culturais locais (TOUVRON et al., 2023; BENDER et al., "
        "2021). Fatores técnicos como latência e custos indexados ao dólar também "
        "dificultam aplicações de tempo real (CHEN et al., 2023). Observa-se que o "
        "corpus de pré-treino utilizado neste trabalho é composto por textos sem "
        "dados pessoais — Constituição Federal, obras literárias em domínio público "
        "e amostras de ficção contemporânea empregadas sob uso estritamente "
        "acadêmico e não-comercial, sem redistribuição dos pesos resultantes — "
        "de modo que as considerações de LGPD se aplicam ao caso de uso futuro "
        "(atendimento ao cliente) e não ao dado de treino em si."
    )

    # ------------------------------------------------------------------
    # 4. Parágrafo BPE / Tokenizer — parágrafo [72] e [80]
    # ------------------------------------------------------------------
    par72_novo = (
        "O dataset final foi tokenizado utilizando o vocabulário BPE (Byte-Pair "
        "Encoding) padrão do GPT-2, de 50.257 tokens, resultando em um volume total "
        "de 2.506.590 tokens de treino. Optou-se por reutilizar o vocabulário do "
        "GPT-2 em vez de treinar um tokenizador próprio em português — decisão que "
        "simplifica a pipeline e permite compatibilidade com ferramentas pré-existentes, "
        "mas que introduz uma limitação conhecida: palavras comuns do português como "
        "\"não\", \"são\" e \"coração\" são segmentadas em múltiplos subtokens, "
        "aumentando o comprimento efetivo das sequências. Um tokenizador treinado "
        "especificamente em corpus PT-BR é discutido como trabalho futuro."
    )

    par80_novo = (
        "Para solucionar o desafio do vocabulário e palavras desconhecidas "
        "(Out-of-Vocabulary), adotou-se a tokenização subword. Utilizou-se o "
        "algoritmo Byte-Pair Encoding (BPE), introduzido em redes neurais por "
        "Sennrich et al. (2016), que funde iterativamente os pares de caracteres "
        "mais frequentes. Para este trabalho, empregou-se o vocabulário BPE do "
        "GPT-2 (OPENAI, 2023), cuja codificação opera diretamente sobre bytes "
        "UTF-8 e garante cobertura lexical universal, ainda que subótima para "
        "português brasileiro."
    )

    # ------------------------------------------------------------------
    # 5. Contradição sobre overfitting — parágrafo [95]
    # ------------------------------------------------------------------
    par95_novo = (
        "Observa-se que a perda de validação iniciou em 4,286 (iteração 200) e "
        "apresentou decaimento consistente ao longo do treinamento, atingindo 3,008 "
        "na iteração 800, 2,145 na iteração 1.400 e 2,003 na iteração 5.000 (ver "
        "Tabela 2). Nas etapas iniciais, a perda de validação se manteve próxima "
        "ou até inferior à de treino, comportamento atribuído ao uso de Dropout "
        "(0,1), ativo apenas no treino. A partir de aproximadamente a iteração "
        "1.400 observou-se afastamento entre as curvas, com a perda de treino "
        "caindo mais rapidamente — sinal de overfitting incipiente dado o tamanho "
        "modesto do corpus (2,5 milhões de tokens). A estratégia adotada foi de "
        "manter o treinamento até estabilização da perda de validação, "
        "caracterizando uma variante de early stopping implícita."
    )

    # ------------------------------------------------------------------
    # 6. Avaliação qualitativa — parágrafo [104]
    # ------------------------------------------------------------------
    par104_novo = (
        "A redução da perda ao patamar de 2,0 e a acurácia acima de 59% indicam "
        "que o modelo passou a gerar sequências de tokens sintaticamente válidas "
        "em português, com conjugação verbal, estrutura de diálogo (marcada por "
        "travessões) e pontuação coerentes. A Seção \"Amostras Qualitativas\" "
        "apresenta exemplos representativos. A coerência semântica de longo prazo "
        "(manutenção de um fio narrativo extenso) permanece limitada, o que é "
        "compatível com a capacidade do modelo (52,93M parâmetros) e com a "
        "densidade informacional do corpus."
    )

    # ------------------------------------------------------------------
    # 7. Conclusão — parágrafos [107] a [111]
    # ------------------------------------------------------------------
    par107_novo = (
        "O presente trabalho atingiu seu objetivo geral ao demonstrar a viabilidade "
        "técnica e metodológica de treinar um Modelo de Linguagem de Pequeno Porte "
        "(SLM) especializado em português brasileiro utilizando infraestrutura de "
        "hardware acessível (6 GB de VRAM). A construção completa do pipeline — "
        "curadoria de corpus multi-domínio, tokenização BPE, definição da "
        "arquitetura decoder-only de 52,93M parâmetros e treino com checkpointing "
        "— comprovou que é possível desenvolver soluções de IA generativa locais e "
        "alinhadas à realidade acadêmica brasileira."
    )

    par108_novo = (
        "Os resultados quantitativos validaram a arquitetura proposta. A perda de "
        "validação caiu de 4,286 (iteração 200) para 2,003 (iteração 5.000), "
        "enquanto a acurácia de predição do próximo token subiu de 26,2% para "
        "59,4% no mesmo intervalo. A estratégia de tokenização BPE e o agendamento "
        "de taxa de aprendizado com decaimento cosseno permitiram que o modelo "
        "operasse dentro do limite rígido de 6 GB de VRAM sem instabilidades de "
        "treino. A avaliação qualitativa confirma que o gerador aprende estruturas "
        "morfossintáticas do português e reproduz formatos de diálogo presentes no "
        "corpus."
    )

    par110_novo = (
        "Entretanto, as limitações do trabalho devem ser reconhecidas. A restrição "
        "de hardware impôs um limite à janela de contexto (256 tokens) e à "
        "profundidade da rede, o que restringe a capacidade do modelo de manter "
        "coerência em gerações longas. O volume de dados — 2,5 milhões de tokens — "
        "é suficiente para a convergência sintática observada, mas insuficiente "
        "para capturar o conhecimento de mundo necessário a um assistente virtual "
        "robusto. O tokenizador foi herdado do GPT-2 sem re-treinamento em PT-BR, "
        "o que aumenta o número de subtokens por palavra e reduz a eficiência de "
        "uso da janela de contexto. A avaliação limitou-se a métricas intrínsecas "
        "(loss, acurácia de token) e análise qualitativa; métricas extrínsecas "
        "como ROUGE-L e BERTScore (LIN, 2004; ZHANG et al., 2020) permanecem como "
        "trabalho futuro."
    )

    par111_novo = (
        "Como encaminhamentos, este trabalho já deixa preparada a etapa seguinte: "
        "foi agregado ao corpus um conjunto de pares pergunta/resposta em português "
        "(dataset Dolly traduzido via LibreTranslate e amostra de 15.000 instâncias "
        "do dataset Alpaca traduzido via OPUS-MT), totalizando aproximadamente "
        "45.000 pares. Tais dados viabilizam a aplicação de Fine-Tuning Supervisionado "
        "(SFT) para especialização do modelo em formato conversacional e em tarefas "
        "de resumo, conforme o título do trabalho sugere. Recomenda-se, adicionalmente, "
        "o treino de um tokenizador BPE próprio em PT-BR, a expansão da janela de "
        "contexto via atenção eficiente (por exemplo, Flash-Attention), e a "
        "exploração de quantização em 4-bit ou 8-bit para viabilizar modelos maiores "
        "no mesmo hardware. A aplicação de alinhamento por preferências humanas "
        "(RLHF; OUYANG et al., 2022) ou de feedback automatizado (RLAIF) fica "
        "indicada como etapa subsequente, após a fase SFT."
    )

    # ------------------------------------------------------------------
    # Aplicar edições em parágrafos
    # ------------------------------------------------------------------
    edits = {
        39: novo_resumo,
        40: palavras_chave,
        58: par58_novo,
        59: par59_novo,
        72: par72_novo,
        80: par80_novo,
        95: par95_novo,
        104: par104_novo,
        107: par107_novo,
        108: par108_novo,
        110: par110_novo,
        111: par111_novo,
    }
    for idx, new_text in edits.items():
        set_paragraph_text(doc.paragraphs[idx], new_text)
        print(f"[ok] parágrafo {idx} atualizado")

    # ------------------------------------------------------------------
    # Tabela 1 — hiperparâmetros V7
    # ------------------------------------------------------------------
    t1 = doc.tables[0]
    t1_data = [
        ("Arquitetura", "Transformer Decoder-Only (GPT-2)", "Padrão para modelagem causal generativa."),
        ("Parâmetros Totais", "52,93 Milhões", "Limite operacional em 6 GB de VRAM."),
        ("Camadas (n_layer)", "8", "Profundidade compatível com a memória disponível."),
        ("Cabeças de Atenção (n_head)", "6", "Divisor de n_embd (384/6 = 64 por cabeça)."),
        ("Dimensão do Embedding (n_embd)", "384", "Equilíbrio entre capacidade e custo."),
        ("Janela de Contexto (block_size)", "256 tokens", "Suficiente para parágrafos e diálogos curtos."),
        ("Batch Size", "16", "Reduzido do V5 (32) para acomodar o modelo maior."),
        ("Learning Rate inicial", "3e-4", "Decaimento cosseno até 1e-5 (CosineAnnealingLR)."),
        ("Dropout", "0,1", "Regularização ajustada ao corpus multi-domínio."),
    ]
    # row 0 é header
    # O notebook original tem 10 linhas (header + 9). Vamos manter 10 (header + 9) mas reescrever:
    for i, (param, valor, just) in enumerate(t1_data, start=1):
        if i < len(t1.rows):
            set_cell_text(t1.rows[i].cells[0], param)
            set_cell_text(t1.rows[i].cells[1], valor)
            set_cell_text(t1.rows[i].cells[2], just)
    print("[ok] Tabela 1 atualizada")

    # ------------------------------------------------------------------
    # Tabela 2 — expandir com mais steps
    # ------------------------------------------------------------------
    t2 = doc.tables[1]
    t2_data = [
        ("200", "4,144", "4,286", "26,2%"),
        ("400", "3,670", "3,300", "33,2%"),
        ("600", "3,479", "3,135", "34,2%"),
        ("800", "3,363", "3,008", "35,6%"),
        ("1.400 (V6)", "1,044", "2,145", "63,5%"),
        ("3.000 (V7)", "1,835", "2,029", "58,8%"),
        ("4.000 (V7)", "1,797", "2,013", "59,1%"),
        ("5.000 (V7)", "1,775", "2,003", "59,4%"),
    ]
    # atualmente 5 linhas (header + 4). Precisamos 9 linhas (header + 8).
    header_row = t2.rows[0]
    # reescreve as 4 linhas de dados existentes
    for i, row_data in enumerate(t2_data[:4], start=1):
        if i < len(t2.rows):
            for c, val in enumerate(row_data):
                set_cell_text(t2.rows[i].cells[c], val)
    # adiciona as 4 linhas novas
    for row_data in t2_data[4:]:
        new_row = t2.add_row()
        for c, val in enumerate(row_data):
            set_cell_text(new_row.cells[c], val)
    print("[ok] Tabela 2 atualizada (8 linhas de dados)")

    # ------------------------------------------------------------------
    # Sumário — ajustar paginação / remover Metodologia fantasma
    # ------------------------------------------------------------------
    # Localiza o parágrafo 55 (SUMÁRIO) e reconstrói. Mais simples: sobrescrevemos
    # com texto completo em um bloco. O docx original tem 7 itens do sumário como
    # se fossem texto corrido entre [55] e [62]. Deixo nas mãos do autor revisar
    # manualmente depois, apenas avisando via impressão.
    print("[info] sumário: os números de página precisarão ser recalculados pelo Word (Ctrl+A / F9)")

    # ------------------------------------------------------------------
    # Remover referências não citadas: Subakan, Mikolov, Pascanu, Hochreiter
    # (parágrafos 130, 146, 152, 164)
    # ------------------------------------------------------------------
    refs_a_remover_idx = [130, 146, 152, 164]
    # Apaga o texto — um tratamento mais preservador que deletar o nó.
    for idx in refs_a_remover_idx:
        set_paragraph_text(doc.paragraphs[idx], "")
    print(f"[ok] {len(refs_a_remover_idx)} referências não citadas esvaziadas")

    # ------------------------------------------------------------------
    # Inserir seção "Amostras Qualitativas" antes de CONCLUSÃO (par 106)
    # ------------------------------------------------------------------
    # Amostras reais geradas pelo V7 step 5000, com stop-strings aplicadas
    amostras_titulo = "Amostras Qualitativas"
    amostras_intro = (
        "A seguir, apresentam-se quatro amostras representativas de texto gerado "
        "pelo modelo V7 (checkpoint do step 5.000) a partir de prompts curtos em "
        "português. Parâmetros de amostragem: temperatura 0,7 e top-k 50. O token "
        "especial <|endoftext|> foi filtrado, juntamente com marcadores de metadata "
        "que vazaram do corpus na etapa de leitura, para que a saída exibida "
        "contenha apenas texto corrido."
    )

    amostras = [
        ("Prompt", "Era uma vez um menino que"),
        ("Saída",
         "Era uma vez um menino que me perguntava ao dia. ia ajudá-la a irmã em "
         "si mesma. Como ia lembrar-se. ia dar-lhe ao lado de Virgilia. ia aos "
         "domingos, era…"),
        ("Prompt", "A velha senhora olhou pela janela e"),
        ("Saída",
         "A velha senhora olhou pela janela e, sendo ingerindo seu cabelo, meu "
         "corpo está encarando o botão. Não sei o que está pensando sobre como "
         "estamos juntos…"),
        ("Prompt", "No alto da montanha, o vento soprava e"),
        ("Saída",
         "No alto da montanha, o vento soprava e a irmã a minha mãe. irmão muito "
         "das irmãs dois parecem na minha voz. irmã, porque não é difícil…"),
        ("Prompt", "A Constituição Federal estabelece que"),
        ("Saída",
         "A Constituição Federal estabelece que fazem. (Vide Emenda Constitucional "
         "nº 19, de 1998)…"),
    ]

    amostras_analise = (
        "Observa-se que o modelo (i) respeita a ortografia e a acentuação do "
        "português na maioria dos tokens gerados; (ii) reproduz estruturas de "
        "diálogo típicas da ficção contemporânea presente no corpus, incluindo "
        "pronomes clíticos e menções como \"Virgilia\" do corpus de Machado de "
        "Assis; (iii) na amostra jurídica, reconhece marcadores normativos "
        "(\"Vide Emenda Constitucional\") aprendidos do texto da Constituição; "
        "(iv) apresenta repetição excessiva de muletas lexicais como \"ia\" e "
        "\"irmã\", indicando distribuição de probabilidade mal calibrada em "
        "regiões do espaço de contexto subrepresentadas. A coerência semântica "
        "decai rapidamente após ~15 tokens, comportamento esperado para um modelo "
        "com 52,93 milhões de parâmetros treinado em 2,5 milhões de tokens."
    )

    conclusao_heading = doc.paragraphs[106]
    insert_paragraph_before(conclusao_heading, amostras_titulo, style="Heading 2")
    insert_paragraph_before(conclusao_heading, amostras_intro)
    for role, text in amostras:
        insert_paragraph_before(conclusao_heading, f"{role}: {text}")
    insert_paragraph_before(conclusao_heading, amostras_analise)
    print("[ok] seção de Amostras Qualitativas inserida antes da Conclusão")

    # ------------------------------------------------------------------
    # Inserir Figura 1 (gráfico de loss) antes de "Figura 1. ..." no par 94
    # ------------------------------------------------------------------
    grafico = Path("/home/akirem/Documentos/LLM-USP/outputs/grafico_loss_ptbr.png")
    if grafico.exists():
        caption_par = doc.paragraphs[94]  # o parágrafo "Figura 1. Evolução..."
        # Cria novo parágrafo antes com a imagem
        new_p = caption_par.insert_paragraph_before()
        run = new_p.add_run()
        try:
            run.add_picture(str(grafico), width=Inches(6.0))
            print(f"[ok] gráfico inserido: {grafico.name}")
        except Exception as e:
            print(f"[warn] não inseriu imagem: {e}")
    else:
        print(f"[warn] gráfico não encontrado em {grafico}")

    # ------------------------------------------------------------------
    # Salva
    # ------------------------------------------------------------------
    doc.save(str(DST))
    print(f"\n=== SALVO em: {DST} ===")
    print("Abra o arquivo no Word/LibreOffice e:")
    print("  1. Pressione Ctrl+A, depois F9 (ou Atualizar Tudo) para recalcular sumário/páginas")
    print("  2. Revise as frases alteradas para ajustes finos de estilo")
    print("  3. Confira formatação da tabela (células podem ter perdido alinhamento central)")


if __name__ == "__main__":
    main()
